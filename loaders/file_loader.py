import os
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
    # UnstructuredWordLoader
)

#from langchain_community.document_loaders import UnstructuredFileLoader #Can use for other file formats also
from docx import Document
from langchain_core.documents import Document as LangchainDocument

def load_docx_file(file_path):
    doc = Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return [
        LangchainDocument(
            page_content=full_text,
            metadata={"source": os.path.basename(file_path)}
        )
    ]


def load_single_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
    elif ext == ".txt":
        loader = TextLoader(file_path, encoding='utf-8')
    elif ext == ".csv":
        loader = CSVLoader(file_path)
    elif ext == ".docx":
        return load_docx_file(file_path)
        # loader = UnstructuredFileLoader(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
    docs = loader.load()
    for doc in docs:
        doc.metadata["source"] = os.path.basename(file_path)
    return docs

def load_documents_from_folder(folder_path):
    all_docs = []
    for file in os.listdir(folder_path):
        full_path = os.path.join(folder_path, file)
        if os.path.isfile(full_path):
            # print(file)
            try:
                all_docs.extend(load_single_file(full_path))
            except Exception as e:
                print(f"⚠️ Failed to load {file}: {e}")
    return all_docs
